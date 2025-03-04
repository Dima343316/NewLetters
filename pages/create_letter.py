import streamlit as st
from datetime import datetime
import os
from docx import Document
import openai
from dotenv import load_dotenv

from minions import letter_to_text
from minions import rag_doc
load_dotenv()

# Класс для работы с моделью GPT (например, Llama)
# Класс для работы с API
class LlamaModel:
    def __init__(self):
        self.api_key = os.getenv('OPENAI_TOKEN_API')  # Чтение ключа API
        self.base_url = os.getenv('BASE_OPENAI_URL')
        self.max_input_tokens = 2048  # Максимум токенов для входного сообщения
        self.max_response_tokens = 2048  # Максимум токенов для ответа

    def count_tokens(self, messages):
        tokens = sum(len(message["content"].split()) for message in messages)
        return tokens

    def get_response(self, message):
        messages = [
            {"role": "user", "content": message}
        ]

        total_tokens = self.count_tokens(messages)
        if total_tokens > self.max_input_tokens:
            return "Ваш вопрос слишком длинный. Пожалуйста, сократите его."

        client = openai.Client(
            api_key=self.api_key,
            base_url=self.base_url,
        )


        response=client.chat.completions.create(
                model="openai/gpt-4o-2024-11-20",  # Укажите модель, например, OpenAI GPT-3
                messages=messages,
                temperature=0.4,  # Степень случайности
                max_tokens=1500,  # Максимальное количество токенов в ответе
                top_p=0.3,  # Использование метода nucleus sampling
                frequency_penalty=0.1,  # Штраф за частоту
                presence_penalty=0.1,  # Штраф за присутствие
            )

        return response.choices[0].message.content


# Класс для обработки шаблона документа и замены значений
class PasteValues:
    def __init__(self, template_path, data):
        self.template_path = template_path
        self.data = data

    def change_values(self, letter_body):
        # Генерируем имя файла с текущей датой и временем
        current_datetime = datetime.now().strftime("%Y%m%d-%H%M%S")
        filename = f"docs/result/result_{current_datetime}.docx"

        # Проверка, существует ли папка для сохранения файла
        os.makedirs(os.path.dirname(filename), exist_ok=True)

        # Открываем шаблон
        doc = Document(self.template_path)

        # Заменяем текст в каждом параграфе
        for p in doc.paragraphs:
            for key, value in self.data.items():
                placeholder = f"${{{key}}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, value)

        # Вставляем сгенерированный текст в body письма
        for p in doc.paragraphs:
            if "${letter_body}" in p.text:
                p.text = p.text.replace("${letter_body}", letter_body)

        # Сохраняем изменённый документ
        doc.save(filename)
        return filename


# Основная страница для ввода данных
def create_letter_page():
    flag_incoming = False
    flag_act = False
    flag_rag = False
    rag_info = ""
    doc_incoming = ""
    st.title("📄 Генерация писем")
    folders = os.listdir(r"db_docs")
    fio_recipient = st.text_input("🔹 ФИО получателя:", placeholder="Введите ФИО...")
    position_recipient = st.text_input("🔹 Должность получателя:", placeholder="Введите должность...")
    laws = st.selectbox("Выберите документ:", folders)
    comments_for_law = st.text_area("📝 Ключевые слова к закону:", placeholder="Укажите ключевые слова к закону...")
    if st.button("📌 Загрузить информацию от базы данных"):
        flag_rag = True
        rag_info = rag_doc.run_gpt_query(comments_for_law, f"db_docs/{laws}")
        st.text_area("Результат анализа входящего письма", rag_info, height=400)
    comments = st.text_area("📝 Комментарии:", placeholder="Укажите комментарии к письму...")
    topic_body = st.text_input("📝 Основная тема письма:", placeholder="Введите тему письма...")
    fio_sender = st.text_input("✍ ФИО отправителя:", placeholder="Введите ваше ФИО...")
    doc_date = st.date_input("📅 Дата создания документа:", value=datetime.today())
    if incoming_letter := st.file_uploader("Загрузите входящее письмо", type=["pdf", "docx", "txt", "odt"]):
        text_of_doc = letter_to_text.extract_text_from_file(incoming_letter)
    if st.button("📌 Загрузить входящее письмо"):
        flag_incoming = True
        llama_model = LlamaModel()
        incoming_gpt = f"""
                    Извлеките из следующего письма следующую информацию:
                    1. **Имя отправителя**: Имя человека, отправившего письмо.
                    2. **Организация отправителя**: Название организации, от которой отправлено письмо.
                    3. **Должность отправителя**: Должность лица, отправившего письмо.
                    4. **Основная часть письма**: Суть письма, исключая приветствия и заключительные фразы.
                    
                    Вот текст письма:{text_of_doc}
                    
                    Пожалуйста, извлеките только указанную информацию в виде списка без добавления дополнительных фраз.
    
                        """
        doc_incoming = llama_model.get_response(incoming_gpt)
        st.text_area("Результат анализа входящего письма", doc_incoming, height=400)
    if st.button("📌 Загрузить акт"):
        flag_act = True
        llama_model = LlamaModel()
        act_gpt = f"""
            Проанализируй текст акта и выдели ключевые моменты. Определи основные факты, выводы, решения и рекомендации, содержащиеся в документе. Укажи, если есть какие-либо важные даты, действия, субъекты, а также любую информацию о последствиях или рекомендациях для дальнейших действий.
            {text_of_doc}
                            """
        doc_incoming = llama_model.get_response(act_gpt)
        st.text_area("Результат анализа акта", doc_incoming, height=400)
    if st.button("📌 Сформировать письмо"):
        # Генерация письма через GPT
        llama_model = LlamaModel()
        gpt_input = f"""
                                **!!Строго!!Не включайте заключительные фразы вроде "С уважением" или других подобных выражений. **
                                    Составьте официальное письмо на основе следующих данных. Письмо должно быть оформлено в строгом деловом стиле, без вступлений и заключений, таких как "С уважением".
                                    В письме должно быть изложено только основное содержание, без дополнительных вежливых фраз. 
                                    Для составления содержания используйте тему письма и комментарии.
                                    1. **Тема письма**: {topic_body}
                                    2. **Комментарии**: {comments}
                                    Не добавляйте вступление или заключение, типа "С уважением", а также избегайте формальностей и излишней вежливости.
                                    """
        if rag_info != "":
            gpt_input = f"""
                                    **!!Строго!!Не включайте заключительные фразы вроде "С уважением" или других подобных выражений. **
                                        Составьте официальное письмо на основе следующих данных. Письмо должно быть оформлено в строгом деловом стиле, без вступлений и заключений, таких как "С уважением".
                                        В письме должно быть изложено только основное содержание, без дополнительных вежливых фраз.
                                        Для составления содержания используйте тему письма и комментарии.
                                        1. **Текст законы который необходимо применить в письме**: {rag_info}
                                        2. **Тема письма**: {topic_body}
                                        3. **Комментарии**: {comments}
                                        Не добавляйте вступление или заключение, типа "С уважением", а также избегайте формальностей и излишней вежливости.
                                    """
        if doc_incoming != "":
            gpt_input = f"""
            **!!Строго!!Не включайте заключительные фразы вроде "С уважением" или других подобных выражений. **
                Составьте официальное письмо на основе следующих данных. Письмо должно быть оформлено в строгом деловом стиле, без вступлений и заключений, таких как "С уважением".
                В письме должно быть изложено только основное содержание, без дополнительных вежливых фраз. 
                Для составления содержания используйте тему письма и комментарии.
                Ответь на входящее письмо {doc_incoming}
                1. **Законы, применяемые в письме**: {laws}
                2. **Тема письма**: {topic_body}
                3. **Комментарии**: {comments}
                Не добавляйте вступление или заключение, типа "С уважением", а также избегайте формальностей и излишней вежливости.
            """
            flag_incoming = False
        if doc_incoming != "" and rag_info != "":
            gpt_input = f"""
                        **!!Строго!!Не включайте заключительные фразы вроде "С уважением" или других подобных выражений. **
                            Составьте официальное письмо на основе следующих данных. Письмо должно быть оформлено в строгом деловом стиле, без вступлений и заключений, таких как "С уважением".
                            В письме должно быть изложено только основное содержание, без дополнительных вежливых фраз. 
                            Для составления содержания используйте тему письма и комментарии.
                            Ответь на входящее письмо {doc_incoming}
                            1. **Текст законы который необходимо применить в письме**: {laws}
                            2. **Тема письма**: {topic_body}
                            3. **Комментарии**: {comments}
                            Не добавляйте вступление или заключение, типа "С уважением", а также избегайте формальностей и излишней вежливости.
                        """
        letter_body = llama_model.get_response(gpt_input)

        # Генерация данных для вставки в шаблон
        data = {
            "fio_recipient": fio_recipient,
            "position_recipient": position_recipient,
            "laws": laws,
            "topic_body": topic_body,
            "fio_sender": fio_sender,
            "doc_date": str(doc_date),
            "letter_body": letter_body,
        }

        # Создание и сохранение файла с новыми данными
        paste_values = PasteValues("docs/sample/docx_sample.docx", data)
        generated_filename = paste_values.change_values(letter_body)

        # Отображаем кнопку для скачивания файла
        with open(generated_filename, "rb") as f:
            st.download_button(
                label="Скачать письмо",
                data=f,
                file_name=generated_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    create_letter_page()
